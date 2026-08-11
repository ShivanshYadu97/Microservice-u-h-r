from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("output/resume/Shivansh_Yadu_Java_Full_Stack_Resume.docx")


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9.2)
    return p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9E2EC", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def clear_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def add_bottom_border(paragraph, color="1F4E79", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.first_child_found_in("w:pBdr")
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(31, 78, 121)
    add_bottom_border(p)


def add_body_paragraph(doc, text, after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.03
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9.4)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.02
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9.15)


def add_role(doc, title, company, location, dates, tech, bullets):
    table = doc.add_table(rows=1, cols=2)
    clear_table_borders(table)
    table.columns[0].width = Inches(4.7)
    table.columns[1].width = Inches(2.3)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    set_cell_text(left, f"{title} | {company}", bold=True)
    p = set_cell_text(right, dates, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for cell in (left, right):
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    if location:
        p_loc = doc.add_paragraph()
        p_loc.paragraph_format.space_after = Pt(1.5)
        run = p_loc.add_run(location)
        run.italic = True
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)
    add_body_paragraph(doc, f"Tech: {tech}", after=1.5)
    for bullet in bullets:
        add_bullet(doc, bullet)


def build_resume():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.3)
    styles["List Bullet"].font.name = "Arial"
    styles["List Bullet"].font.size = Pt(9.15)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(1)
    run = name.add_run("SHIVANSH YADU")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("Java Full Stack Developer | Spring Boot | Angular | REST APIs")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10.2)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(4)
    run = contact.add_run("Mumbai, India | +91 9340348901 | shivanshyadu9340@gmail.com")
    run.font.name = "Arial"
    run.font.size = Pt(9.2)
    run.font.color.rgb = RGBColor(70, 70, 70)

    add_section_heading(doc, "Profile Summary")
    add_body_paragraph(
        doc,
        "Java Full Stack Developer with 2.3 years of experience building web applications using "
        "Java 8, Spring Boot, REST APIs, Hibernate/JPA, Angular, TypeScript, MySQL and PostgreSQL. "
        "Hands-on with authentication and authorization flows using Spring Security and JWT, MVC "
        "architecture, API documentation with Swagger, Git-based collaboration, and Maven-driven "
        "project builds. Strong fit for Java Full Stack, Spring Boot Developer and Java Angular roles.",
        after=2,
    )

    add_section_heading(doc, "Technical Skills")
    skills = [
        ("Backend", "Java, Java 8, Spring Boot, RESTful APIs, Spring MVC, Spring Security, JWT, Hibernate/JPA, JDBC, HQL"),
        ("Frontend", "Angular, TypeScript, RxJS, JSP, jQuery, Bootstrap, Tailwind CSS, HTML templates"),
        ("Databases", "MySQL, PostgreSQL, Oracle"),
        ("Tools", "Git/GitHub, Maven, Postman, Swagger UI, IntelliJ IDEA, STS, Eclipse, VS Code, MySQL Workbench"),
        ("Other", "OOPs, Collections, Lambda, Streams, Design Patterns, Logging, AWS basics, Checkmarx"),
    ]
    table = doc.add_table(rows=len(skills), cols=2)
    table.autofit = False
    for row, (label, text) in zip(table.rows, skills):
        row.cells[0].width = Inches(1.15)
        row.cells[1].width = Inches(6.25)
        set_cell_shading(row.cells[0], "EAF2F8")
        set_cell_border(row.cells[0])
        set_cell_border(row.cells[1])
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_text(row.cells[1], text)

    add_section_heading(doc, "Professional Experience")
    add_role(
        doc,
        "Java Developer",
        "Housing.com Pvt. Ltd.",
        "",
        "Feb 2025 - Present",
        "Java, Spring Boot, Angular, TypeScript, Hibernate/JPA, MySQL, REST APIs, RestTemplate, MVC",
        [
            "Built and enhanced modules for a real-estate web application covering property listing, property search, user authentication and role-based access for owners, buyers and admin users.",
            "Developed Spring Boot REST APIs and Angular screens for property filters and admin workflows, improving backend task handling and day-to-day management operations.",
            "Worked with Hibernate/JPA and MySQL for persistence, query handling and data retrieval across property and user-management flows.",
            "Integrated internal and external REST APIs using RestTemplate and validated endpoints through Postman and Swagger UI.",
            "Collaborated in a 7-member team using Git/GitHub, Maven and IDE debugging to deliver feature fixes and application improvements.",
        ],
    )
    add_role(
        doc,
        "Java Full Stack Developer",
        "TransUnion CIBIL",
        "Mumbai",
        "May 2024 - Jan 2025",
        "Java, Spring Boot, Angular, TypeScript, Hibernate/JPA, PostgreSQL, Bootstrap, jQuery, MVC",
        [
            "Developed SSO integration features that consumed multiple APIs and supported user verification across five databases before enabling single sign-on access.",
            "Implemented REST controller, service and repository layers in Spring Boot with Hibernate/JPA and PostgreSQL for verification and user-data workflows.",
            "Built Angular, TypeScript, Bootstrap and jQuery-based UI components for registration, login verification and user-facing SSO flows.",
            "Supported authentication and authorization implementation using Spring Security and JWT-based access control.",
            "Prepared and tested API contracts with Swagger UI and Postman while coordinating with a 10-member engineering team.",
        ],
    )

    add_section_heading(doc, "Core Responsibilities")
    for item in [
        "Designed and implemented REST APIs with request validation, layered Spring Boot architecture and database integration.",
        "Worked on database operations using Hibernate/JPA, JDBC, HQL, MySQL and PostgreSQL.",
        "Applied Java 8 features such as lambdas, streams and collections while writing maintainable application logic.",
        "Handled debugging, issue fixing, dependency management and Git-based version control in active development environments.",
    ]:
        add_bullet(doc, item)

    add_section_heading(doc, "Education")
    table = doc.add_table(rows=1, cols=2)
    clear_table_borders(table)
    set_cell_text(table.cell(0, 0), "B.Tech - Chhattisgarh Swami Vivekanand Technical University, Bhilai", bold=True)
    p = set_cell_text(table.cell(0, 1), "CGPA: 8.1", bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    add_section_heading(doc, "Resume Keywords")
    add_body_paragraph(
        doc,
        "Java Full Stack Developer, Java Developer, Spring Boot Developer, Java Angular Developer, "
        "REST API Developer, Hibernate JPA, MySQL, PostgreSQL, Spring Security, JWT, Angular, TypeScript.",
        after=0,
    )

    doc.save(OUT)


if __name__ == "__main__":
    OUT.parent.mkdir(parents=True, exist_ok=True)
    build_resume()
    print(OUT.resolve())
