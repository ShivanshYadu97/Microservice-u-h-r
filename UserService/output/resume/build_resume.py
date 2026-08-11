from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "output/resume/Shivansh_Yadu_Java_Full_Stack_Resume.docx"


BLUE = RGBColor(31, 78, 121)
TEXT = RGBColor(32, 32, 32)
MUTED = RGBColor(85, 85, 85)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_in):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_text(paragraph, text, size=10.5, bold=False, color=TEXT):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return run


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    set_run_font(r, size=11, bold=True, color=BLUE)
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9E2F3")
    border.append(bottom)
    p_pr.append(border)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    add_text(p, text, size=9.6)
    return p


def add_role(doc, title, company, location, dates, project, tech, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    add_text(p, title, size=10.3, bold=True, color=TEXT)
    add_text(p, f" | {company}", size=10.2, color=TEXT)
    add_text(p, f" | {location}", size=10.0, color=MUTED)
    add_text(p, f" | {dates}", size=10.0, bold=True, color=TEXT)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    add_text(p2, "Project: ", size=9.5, bold=True)
    add_text(p2, project, size=9.5)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(2)
    add_text(p3, "Tech Stack: ", size=9.5, bold=True)
    add_text(p3, tech, size=9.5)

    for bullet in bullets:
        add_bullet(doc, bullet)


def add_skill_table(doc):
    rows = [
        ("Backend", "Java, Java 8, Spring Boot, Spring MVC, RESTful APIs, Hibernate/JPA, JDBC, Servlets, JSP"),
        ("Frontend", "Angular, TypeScript, RxJS, jQuery, HTML, Bootstrap, Tailwind CSS"),
        ("Database", "MySQL, PostgreSQL, Oracle, HQL"),
        ("Security & Cloud", "Spring Security, JWT authentication, AWS basics, Checkmarx"),
        ("Tools", "Git/GitHub, Maven, Postman, Swagger UI, IntelliJ IDEA, Eclipse, STS, VS Code, MySQL Workbench"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (label, value) in enumerate(rows):
        cells = table.rows[i].cells
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
        set_cell_width(cells[0], 1.35)
        set_cell_width(cells[1], 5.15)
        set_cell_shading(cells[0], "EEF3F8")
        p0 = cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        add_text(p0, label, size=9.4, bold=True, color=BLUE)
        p1 = cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        add_text(p1, value, size=9.4)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    for style_name in ("List Bullet",):
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(9.6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1)
    add_text(title, "SHIVANSH YADU", size=18, bold=True, color=BLUE)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(4)
    add_text(contact, "Mumbai, India | +91 9340348901 | shivanshyadu9340@gmail.com", size=9.5, color=MUTED)

    headline = doc.add_paragraph()
    headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headline.paragraph_format.space_after = Pt(7)
    add_text(
        headline,
        "Java Full Stack Developer | Spring Boot | Angular | REST APIs | Hibernate/JPA | MySQL/PostgreSQL",
        size=10.2,
        bold=True,
        color=TEXT,
    )

    add_heading(doc, "Profile Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_text(
        p,
        "Java Full Stack Developer with 2.3 years of experience building web applications using Java, Spring Boot, Angular, TypeScript, Hibernate/JPA, RESTful APIs, and relational databases. Strong understanding of OOP, Java 8 features, layered MVC architecture, API integration, authentication/authorization using Spring Security and JWT, and database operations with MySQL and PostgreSQL. Comfortable working with Git, Maven, Postman, Swagger, debugging tools, and deployment basics on AWS.",
        size=9.7,
    )

    add_heading(doc, "Technical Skills")
    add_skill_table(doc)

    add_heading(doc, "Work Experience")
    add_role(
        doc,
        "Java Developer",
        "Housing.com Pvt Ltd",
        "Mumbai",
        "Feb 2025 - Present",
        "Zillow Scraper / Real Estate Web Application",
        "Java, Spring Boot, Angular, TypeScript, Hibernate/JPA, MySQL, REST Template, MVC Architecture",
        [
            "Developed Spring Boot and Angular features for property listings, search, user authentication, and role-based access for owners and buyers.",
            "Integrated REST APIs and backend services to support property data retrieval, filtering, and management workflows.",
            "Built admin-side listing and filter features to make property management faster and easier for internal users.",
            "Used Hibernate/JPA with MySQL for persistence, query handling, and database operations.",
            "Collaborated with a 7-member team across requirement understanding, feature development, debugging, and delivery.",
        ],
    )

    add_role(
        doc,
        "Java Full Stack Developer",
        "TransUnion CIBIL",
        "Mumbai",
        "May 2024 - Jan 2025",
        "SSO Integration Application",
        "Java, Spring Boot, Angular, TypeScript, Hibernate/JPA, PostgreSQL, Bootstrap, jQuery, MVC Architecture",
        [
            "Developed and launched an SSO integration application that consumed multiple APIs and retrieved user data across five databases.",
            "Implemented backend flows for new user registration, existing user verification, and single sign-on enablement.",
            "Built RESTful APIs using Spring Boot and Hibernate/JPA, following controller, service, and repository layering.",
            "Worked on Angular and TypeScript UI components with Bootstrap and jQuery for user-facing verification flows.",
            "Supported authentication and authorization implementation using Spring Security and JWT tokens.",
        ],
    )

    add_heading(doc, "Key Responsibilities")
    for item in [
        "Designed and implemented REST controllers, service-layer logic, repository operations, and validation flows in Spring Boot.",
        "Worked with Java 8 features including lambdas, streams, collections, string handling, and OOP principles.",
        "Prepared API documentation and testing flows using Swagger UI and Postman.",
        "Used Maven for dependency management and Git/GitHub for version control.",
        "Debugged defects using IntelliJ IDEA, Eclipse, STS, and application logs.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Education")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    add_text(p, "B.Tech", size=10.0, bold=True)
    add_text(p, " | Chhattisgarh Swami Vivekananda Technical University, Bhilai | CGPA: 8.1", size=10.0)

    doc.save(OUT)


if __name__ == "__main__":
    build()
